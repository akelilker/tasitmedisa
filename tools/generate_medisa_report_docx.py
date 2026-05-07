from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "outputs" / "medisa-tanitim-raporu-guncel-2026-05-07.docx"
LOGO = ROOT / "icon" / "medlogo.png"


ACCENT = RGBColor(8, 13, 22)
ACCENT_SOFT = RGBColor(52, 73, 94)
MUTED = RGBColor(96, 108, 118)
LIGHT = RGBColor(236, 242, 247)
BORDER = "D7E0E8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D7E0E8", size="8"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_font(run, name="Aptos", size=Pt(11), bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def style_normal(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.08

    for style_name, size in [("Title", 24), ("Heading 1", 15), ("Heading 2", 12)]:
        style = styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO.exists():
        p.add_run().add_picture(str(LOGO), width=Inches(1.8))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("MEDISA Taşıt Yönetim Sistemi")
    set_font(r, size=Pt(24), bold=True, color=ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("Güncel Tanıtım ve Durum Raporu")
    set_font(r, size=Pt(14), bold=True, color=ACCENT_SOFT)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(18)
    r = meta.add_run("Rapor Tarihi: 7 Mayıs 2026")
    set_font(r, size=Pt(10.5), color=MUTED)

    box = doc.add_table(rows=3, cols=2)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    widths = [Cm(5.2), Cm(10.2)]
    for row in box.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
    items = [
        ("Proje Tipi", "Kurumsal taşıt kayıt, takip, raporlama ve kullanıcı yönetim sistemi"),
        ("Teknik Yapı", "PHP 8.2 + Vanilla HTML/CSS/JavaScript + JSON veri saklama"),
        ("Güncel Kapsam", "Ana yönetim paneli, kullanıcı paneli, admin raporlama, belge yönetimi, PWA"),
    ]
    for i, (label, value) in enumerate(items):
        left, right = box.rows[i].cells
        set_cell_shading(left, "EDF3F7")
        set_cell_shading(right, "FFFFFF")
        p1 = left.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(label)
        set_font(r1, size=Pt(10.5), bold=True, color=ACCENT)
        p2 = right.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r2 = p2.add_run(value)
        set_font(r2, size=Pt(10.5))

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    note.paragraph_format.left_indent = Cm(0.3)
    note.paragraph_format.right_indent = Cm(0.3)
    r = note.add_run(
        "Bu doküman, erken dönem taslak tanıtım yerine projenin mevcut üretim olgunluğunu "
        "yansıtmak amacıyla güncel kod tabanı ve aktif modüller baz alınarak hazırlanmıştır."
    )
    set_font(r, size=Pt(10.5), color=ACCENT_SOFT)

    doc.add_page_break()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_font(r, size=Pt(15 if level == 1 else 12), bold=True, color=ACCENT)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="Normal")
        p.style = doc.styles["Normal"]
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run("• ")
        set_font(r1, size=Pt(10.5), bold=True, color=ACCENT)
        r2 = p.add_run(item)
        set_font(r2, size=Pt(10.5))


def add_paragraph(doc, text, color=None, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, size=Pt(10.5), bold=True, color=ACCENT)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2, size=Pt(10.5), color=color)
    else:
        r = p.add_run(text)
        set_font(r, size=Pt(10.5), color=color)


def add_two_col_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.rows[0].cells[0].width = Cm(5.5)
    table.rows[0].cells[1].width = Cm(11.8)
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Başlık", "Açıklama"]):
        set_cell_shading(cell, "DCE8F2")
        set_cell_border(cell, color="B7C8D6", size="10")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_font(r, size=Pt(10.5), bold=True, color=ACCENT)
    for left_text, right_text in rows:
        left, right = table.add_row().cells
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
        lp = left.paragraphs[0]
        lr = lp.add_run(left_text)
        set_font(lr, size=Pt(10.2), bold=True, color=ACCENT_SOFT)
        rp = right.paragraphs[0]
        rr = rp.add_run(right_text)
        set_font(rr, size=Pt(10.2))
    return table


def add_module_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(3.2), Cm(3.8), Cm(5.4), Cm(4.9)]
    headers = ["Modül", "Giriş Noktası", "Amaç", "Öne Çıkan Yetkinlik"]
    for idx, cell in enumerate(table.rows[0].cells):
        cell.width = widths[idx]
        set_cell_shading(cell, "DCE8F2")
        set_cell_border(cell, color="B7C8D6", size="10")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(headers[idx])
        set_font(r, size=Pt(10), bold=True, color=ACCENT)
    rows = [
        ("Ana Panel", "index.html", "Taşıt kayıt, listeleme ve ana operasyon ekranı", "Kayıt işlemleri, raporlar, ayarlar ve yazdırma akışları"),
        ("Kullanıcı Girişi", "driver/index.html", "Saha kullanıcısının sisteme güvenli erişimi", "Beni hatırla, oturum akışı, PWA uyumu"),
        ("Kullanıcı Paneli", "driver/dashboard.html", "Zimmetli taşıt ve kullanıcı talepleri", "Şifre değiştirme, geçmiş, talep/şikayet/öneri, bildirimler"),
        ("Admin Raporlama", "admin/driver-report.html", "Yönetimsel analiz ve aylık takip", "Aylık KM izleme, kullanıcı raporları, filtreleme, Excel dışa aktarma"),
        ("Belge Yönetimi", "ruhsat.php / upload_ruhsat.php", "Araç belgelerinin güvenli saklanması", "Ruhsat yükleme, önizleme, inline görüntüleme, yetkili erişim"),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for idx, text in enumerate(row_data):
            row[idx].width = widths[idx]
            row[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(row[idx])
            p = row[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx != 1 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_font(r, size=Pt(9.7))
    return table


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MEDISA Taşıt Yönetim Sistemi  •  Güncel Tanıtım Raporu  •  7 Mayıs 2026")
    set_font(r, size=Pt(8.5), color=MUTED)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    style_normal(doc)
    add_footer(doc)
    add_cover(doc)

    add_heading(doc, "1. Yönetici Özeti", 1)
    add_paragraph(
        doc,
        "MEDISA Taşıt Yönetim Sistemi, başlangıçtaki temel araç kayıt yaklaşımından çıkarak "
        "bugün kurumsal ölçekte kullanılabilecek daha kapsamlı bir operasyon paneline dönüşmüştür. "
        "Sistem artık yalnızca araç ekleme ve listeleme yapmamakta; kullanıcı rolleri, raporlama, "
        "belge yönetimi, yedekleme ve mobil kullanım gereksinimlerini birlikte karşılamaktadır.",
    )
    add_bullets(doc, [
        "Ana uygulama, kullanıcı paneli ve admin raporlama modülleri tek veri yapısı üzerinde birlikte çalışır.",
        "Sunucu tarafında JSON tabanlı veri saklama sürdürülür; yazma işlemlerinde snapshot yedek ve geri yükleme desteği bulunur.",
        "Kullanıcı deneyimi tarafında PWA kabiliyeti, mobil uyum ve offline dayanıklılığı güçlendirilmiştir.",
        "Ruhsat gibi araç belgeleri için yükleme, önizleme ve erişim yetkisi kontrollü biçimde yönetilir.",
    ])

    add_heading(doc, "2. Teknik Mimari", 1)
    add_two_col_table(doc, [
        ("Uygulama Katmanı", "PHP 8.2 backend uçları ile vanilla HTML, CSS ve JavaScript tabanlı istemci mimarisi"),
        ("Veri Saklama", "Ana veri deposu olarak data/data.json kullanılır; ek dosya alanları ve belge klasörleri desteklenir"),
        ("Sunucu Yazma Disiplini", "Kayıtlar save.php ve core.php içindeki merkezi akışlar üzerinden atomik yazım mantığıyla saklanır"),
        ("Yedekleme Kurgusu", "data.json.backup ve zaman damgalı snapshot dosyaları ile veri kaybı riski azaltılır"),
        ("Dağıtım Yapısı", "Apache + mod_rewrite ile çalışır; cPanel Git veya GitHub Actions FTP deploy akışına uygundur"),
    ])

    add_heading(doc, "3. Güncel Modüller", 1)
    add_module_table(doc)

    add_heading(doc, "4. Operasyonel Güçlü Yanlar", 1)
    add_bullets(doc, [
        "Rol bazlı yetki sınırları sayesinde genel yönetici ve sınırlı kapsamlı kullanıcı akışları ayrıştırılmıştır.",
        "Bildirim okuma durumu, kullanıcı kapsamı ve şube kırılımı gibi yardımcı durumlar veri modeli içinde izlenebilmektedir.",
        "Aylık kilometre takibi ve kullanıcı raporları yönetim tarafında ölçülebilir görünürlük sağlar.",
        "Ruhsat yükleme ve görüntüleme akışı, belgeyi dosya sistemi üzerinde tutarken veriye kontrollü bağlantı kurar.",
        "Service Worker ve manifest kurgusu sayesinde uygulama PWA olarak kurulabilir ve temel kabuk offline dayanıklılığı kazanır.",
    ])

    add_heading(doc, "5. Veri Güvenliği ve Süreklilik", 1)
    add_paragraph(
        doc,
        "Sistemde veritabanı yerine düz dosya mimarisi kullanılmasına rağmen kayıt güvenliği için önemli koruma adımları eklenmiştir. "
        "Yazma işleminden önce önceki sürüm yedeklenmekte, ayrıca zaman damgalı snapshot dizini tutulmaktadır. "
        "Geri yükleme tarafında restore.php ile son güvenli kopyaya dönüş imkanı sağlanmaktadır."
    )
    add_bullets(doc, [
        "Atomik yazım yaklaşımı yarım kalmış kayıt riskini azaltır.",
        "Ana yedek ve snapshot mantığı operasyonel hata anlarında geri dönüş kolaylığı sağlar.",
        "Belge dosyaları için doğrudan data dizinine erişim .htaccess kurallarıyla sınırlanmıştır.",
    ])

    add_heading(doc, "6. Kullanıcı Deneyimi ve Mobil Uyum", 1)
    add_paragraph(
        doc,
        "Güncel yapı masaüstü kullanımın yanında mobil/PWA senaryolarına da odaklanmaktadır. "
        "Kullanıcı paneli, giriş ekranı ve admin raporlama sayfalarında responsive düzen, PWA kurulum desteği "
        "ve mobilde daha akıcı kullanım için özel dokunuşlar uygulanmıştır."
    )
    add_bullets(doc, [
        "PWA manifest ve service worker ile uygulama cihaza kurulabilir yapıdadır.",
        "Kritik uygulama kabuğu dosyaları önbelleğe alınarak temel erişim süreleri iyileştirilmiştir.",
        "Kullanıcı panelinde işlem geçmişi, bildirim alanı ve hızlı aksiyon butonları mobil akışa uygun düzenlenmiştir.",
    ])

    add_heading(doc, "7. Mevcut Sınırlamalar", 1)
    add_bullets(doc, [
        "Veri katmanı hâlâ tek JSON dosyası üzerinde çalıştığı için çok yüksek eşzamanlılık senaryolarında doğal ölçek sınırı bulunur.",
        "Otomatik test ve lint altyapısı sınırlıdır; doğrulama halen büyük ölçüde manuel smoke kontrol ile desteklenmektedir.",
        "Kurumsal büyüme devam ederse kullanıcı, belge ve hareket kayıtları için daha güçlü bir kalıcı veri katmanı ihtiyacı oluşacaktır.",
    ])

    add_heading(doc, "8. Kısa Vadeli Geliştirme Önerileri", 1)
    add_bullets(doc, [
        "İlk öncelikte kritik kullanıcı akışları için tarayıcı tabanlı smoke otomasyonunu genişletmek.",
        "Belge yönetimi ve hareket kayıtları için denetlenebilir işlem loglarını daha görünür hale getirmek.",
        "JSON veri katmanını bozmadan, ileride veritabanına geçişi kolaylaştıracak servis sınırlarını netleştirmek.",
        "Yönetim raporlarında dönemsel karşılaştırma, şube bazlı trend ve uyarı metriklerini artırmak.",
    ])

    add_heading(doc, "9. Sonuç", 1)
    add_paragraph(
        doc,
        "MEDISA Taşıt Yönetim Sistemi mevcut haliyle ilk taslak dönemine göre belirgin biçimde olgunlaşmış, "
        "operasyonel süreçleri daha bütünlüklü ele alan bir kurumsal çözüm seviyesine yaklaşmıştır. "
        "Özellikle kullanıcı paneli, admin raporlama, belge akışları ve veri güvenliği yönündeki geliştirmeler, "
        "projenin sadece bir kayıt ekranı değil, sürdürülebilir bir iç operasyon platformu olma yönünde ilerlediğini göstermektedir."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
